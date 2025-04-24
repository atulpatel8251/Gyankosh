import streamlit as st
from docx import Document
import pandas as pd
import io

def extract_table_data(table) -> list:
    """
    Extract data from a table into a list of rows.
    
    Args:
        table: docx Table object
    
    Returns:
        list: List of rows, where each row is a list of cell values
    """
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            # Get text content and clean it
            text = cell.text.strip()
            row_data.append(text)
        data.append(row_data)
    return data

def main():
    st.title("DOCX Table Extractor")
    st.write("""
    Upload a DOCX file to extract and view its table contents.
    Each table will be displayed separately and can be downloaded as CSV.
    """)
    
    uploaded_file = st.file_uploader("Choose a DOCX file", type="docx")
    
    if uploaded_file:
        try:
            # Read the document
            doc_bytes = io.BytesIO(uploaded_file.read())
            doc = Document(doc_bytes)
            
            # Check if document has tables
            if len(doc.tables) == 0:
                st.warning("No tables found in the document!")
            else:
                st.success(f"Found {len(doc.tables)} tables in the document")
                
                # Process each table
                for i, table in enumerate(doc.tables, 1):
                    st.subheader(f"Table {i}")
                    
                    # Extract table data
                    data = extract_table_data(table)
                    
                    # Convert to DataFrame
                    if data:
                        # Use first row as headers
                        headers = data[0]
                        df = pd.DataFrame(data[1:], columns=headers)
                        
                        # Display the table
                        st.dataframe(df)
                        
                        # Convert DataFrame to CSV
                        csv = df.to_csv(index=False)
                        
                        # Download button
                        st.download_button(
                            label=f"Download Table {i} as CSV",
                            data=csv,
                            file_name=f"table_{i}.csv",
                            mime="text/csv"
                        )
                        
                        # Show raw data option
                        with st.expander(f"View Raw Data for Table {i}"):
                            st.write(data)
                    else:
                        st.warning(f"Table {i} is empty")
                    
                    st.markdown("---")  # Add separator between tables
                    
        except Exception as e:
            st.error(f"An error occurred while processing the document: {str(e)}")
            st.exception(e)

if __name__ == "__main__":
    main()